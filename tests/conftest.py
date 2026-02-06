"""Shared fixtures for CreateSlide tests."""

from __future__ import annotations

import sys
from pathlib import Path
from unittest.mock import MagicMock

import pytest

# Ensure project root is on path
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

# ── Environment isolation ───────────────────────────────────────────────
# Prevent tests from touching real API keys / services


@pytest.fixture(autouse=True)
def _isolated_env(monkeypatch, tmp_path):
    """Isolate every test from real env vars and filesystem side effects."""
    monkeypatch.setenv("GOOGLE_API_KEY", "")
    monkeypatch.setenv("OPENAI_API_KEY", "")
    monkeypatch.setenv("OLLAMA_BASE_URL", "http://localhost:11444/v1")
    monkeypatch.setenv("OLLAMA_API_KEY", "test-key")
    monkeypatch.setenv("DEFAULT_PROVIDER", "ollama")
    # Use a temp cancel signal file so tests don't pollute the workspace
    monkeypatch.setenv("CANCEL_SIGNAL_FILE", str(tmp_path / "cancel.flag"))

    # Clear the cached settings singleton so each test picks up monkeypatched env
    from app.config import get_settings

    get_settings.cache_clear()
    yield
    get_settings.cache_clear()


# ── Reusable fixtures ──────────────────────────────────────────────────


@pytest.fixture
def sample_pdf_bytes():
    """Minimal valid PDF bytes for testing."""
    import io

    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.drawString(72, 700, "This is a test document about Artificial Intelligence.")
    c.drawString(72, 680, "AI has transformed many industries including healthcare.")
    c.showPage()
    c.save()
    return buf.getvalue()


@pytest.fixture
def sample_docx_bytes():
    """Minimal valid DOCX bytes for testing."""
    import io

    import docx

    doc = docx.Document()
    doc.add_paragraph("Test document about machine learning.")
    doc.add_paragraph("Neural networks are powerful tools.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_epub_bytes():
    """Minimal valid EPUB bytes for testing."""
    import io
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # EPUB requires specific structure
        zf.writestr("mimetype", "application/epub+zip")
        zf.writestr(
            "META-INF/container.xml",
            '<?xml version="1.0"?>'
            '<container xmlns="urn:oasis:names:tc:opendocument:xmlns:container" version="1.0">'
            "<rootfiles>"
            '<rootfile full-path="content.opf" media-type="application/oebps-package+xml"/>'
            "</rootfiles>"
            "</container>",
        )
        zf.writestr(
            "content.opf",
            '<?xml version="1.0"?>'
            '<package xmlns="http://www.idpf.org/2007/opf" version="2.0" unique-identifier="id">'
            '<metadata xmlns:dc="http://purl.org/dc/elements/1.1/">'
            "<dc:title>Test Book</dc:title>"
            '<dc:identifier id="id">test-id</dc:identifier>'
            "</metadata>"
            "<manifest>"
            '<item id="ch1" href="chapter1.xhtml" media-type="application/xhtml+xml"/>'
            "</manifest>"
            "<spine>"
            '<itemref idref="ch1"/>'
            "</spine>"
            "</package>",
        )
        zf.writestr(
            "chapter1.xhtml",
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<html xmlns="http://www.w3.org/1999/xhtml">'
            "<head><title>Chapter 1</title></head>"
            "<body>"
            "<h1>Chapter 1: Introduction to Artificial Intelligence</h1>"
            "<p>Artificial intelligence is a branch of computer science that aims to create "
            "intelligent machines that can perform tasks that typically require human intelligence. "
            "This includes learning, reasoning, problem-solving, and understanding natural language.</p>"
            "</body>"
            "</html>",
        )
    return buf.getvalue()


@pytest.fixture
def mock_llm_response():
    """Factory fixture that returns a mock LLM provider with a preset response."""

    def _factory(response_text: str = '{"title": "Test", "slides": []}', model: str = "test-model"):
        provider = MagicMock()
        provider.generate.return_value = (response_text, model)
        provider.name = "mock"
        return provider

    return _factory
