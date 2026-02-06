"""Tests for app.services.document — text extraction from files."""

from __future__ import annotations

import io
import tempfile

import pytest

from app.services.document import (
    extract_text_from_docx,
    extract_text_from_epub,
    extract_text_from_pdf,
    load_document,
)


class TestLoadDocument:
    """Test the document loader dispatcher."""

    def test_unsupported_mime_raises(self):
        with pytest.raises(ValueError, match="không được hỗ trợ"):
            load_document(b"data", "text/plain")

    def test_empty_bytes_raises(self):
        with pytest.raises(ValueError, match="rỗng"):
            load_document(b"", "application/pdf")

    def test_load_pdf(self, sample_pdf_bytes):
        text = load_document(sample_pdf_bytes, "application/pdf")
        assert "Artificial Intelligence" in text or "test document" in text.lower()

    def test_load_docx(self, sample_docx_bytes):
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        text = load_document(sample_docx_bytes, mime)
        assert "machine learning" in text.lower()

    def test_load_epub(self, sample_epub_bytes):
        text = load_document(sample_epub_bytes, "application/epub+zip")
        assert "artificial intelligence" in text.lower() or len(text) > 10

    def test_all_supported_mimes(self):
        """Verify all three MIME types are recognized (no crash on dispatch)."""
        supported = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/epub+zip",
        ]
        for mime in supported:
            # These will fail at extraction but should not fail at dispatch
            try:
                load_document(b"dummy", mime)
            except ValueError as e:
                # Only extraction errors, not "not supported" errors
                assert "không được hỗ trợ" not in str(e)
            except Exception:
                pass  # extraction failure with bad bytes is expected


class TestExtractPdf:
    """Test PDF text extraction."""

    def test_valid_pdf(self, sample_pdf_bytes):
        text = extract_text_from_pdf(sample_pdf_bytes)
        assert len(text) > 10
        assert "Artificial Intelligence" in text or "test document" in text.lower()

    def test_invalid_pdf_raises(self):
        with pytest.raises(Exception):
            extract_text_from_pdf(b"not a pdf file")

    def test_multi_page_pdf(self):
        """Multi-page PDFs should concatenate all pages."""
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        c.drawString(72, 700, "Page one content about science.")
        c.showPage()
        c.drawString(72, 700, "Page two content about technology.")
        c.showPage()
        c.save()
        text = extract_text_from_pdf(buf.getvalue())
        assert "science" in text.lower()
        assert "technology" in text.lower()


class TestExtractDocx:
    """Test DOCX text extraction."""

    def test_valid_docx(self, sample_docx_bytes):
        text = extract_text_from_docx(sample_docx_bytes)
        assert "machine learning" in text.lower()

    def test_invalid_docx_raises(self):
        with pytest.raises(Exception):
            extract_text_from_docx(b"not a docx file")

    def test_multi_paragraph_docx(self):
        """DOCX with multiple paragraphs should join them."""
        import docx

        doc = docx.Document()
        doc.add_paragraph("First paragraph about data.")
        doc.add_paragraph("Second paragraph about models.")
        doc.add_paragraph("Third paragraph about training.")
        buf = io.BytesIO()
        doc.save(buf)
        text = extract_text_from_docx(buf.getvalue())
        assert "data" in text.lower()
        assert "models" in text.lower()
        assert "training" in text.lower()


class TestExtractEpub:
    """Test EPUB text extraction."""

    def test_valid_epub(self, sample_epub_bytes):
        text = extract_text_from_epub(sample_epub_bytes)
        assert len(text) > 10

    def test_invalid_epub_raises(self):
        with pytest.raises(Exception):
            extract_text_from_epub(b"not an epub file")

    def test_epub_temp_file_cleanup(self, sample_epub_bytes):
        """After extraction, temp file should be cleaned up."""
        import os

        # Count files in temp dir before
        temp_dir = tempfile.gettempdir()
        epub_files_before = [f for f in os.listdir(temp_dir) if f.endswith(".epub")]
        extract_text_from_epub(sample_epub_bytes)
        epub_files_after = [f for f in os.listdir(temp_dir) if f.endswith(".epub")]
        # Should not leave extra epub files
        assert len(epub_files_after) <= len(epub_files_before)
